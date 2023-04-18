from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx import Document
from docx.shared import Pt
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from bs4 import BeautifulSoup
import json
import re
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities

# Load JSON file containing LinkedIn login information
with open('linkedin_login.json', 'r') as f:
    login_info = json.load(f)

# Retrieve login information
username = login_info['username']
password = login_info['password']

# path to chrome e.g 'C:\Program Files\Google\Chrome\Application\chrome.exe'
chrome_path = 'C:\Program Files\Google\Chrome\Application\chrome.exe'
# path to driver e.g 'C:\path\to\chromedriver.exe'
driver_path = 'C:\Program Files\Google\Chrome\Application\chromedriver.exe'

# Define Chrome options
chrome_options = Options()
chrome_options.binary_location = chrome_path
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-extensions')
chrome_options.add_argument('--headless')
chrome_options.add_argument('--disable-gpu')
chrome_options.add_argument('--disable-dev-shm-usage')

caps = DesiredCapabilities.CHROME
caps['goog:loggingPrefs'] = {'performance': 'ALL'}

service = Service(driver_path)
driver = webdriver.Chrome(
    service=service, options=chrome_options, desired_capabilities=caps)

# Navigate to your LinkedIn profile
driver.get('https://www.linkedin.com/in/adryan-serage')

# Wait for page to load
page_loaded = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.XPATH, '//*[@id="ember10"]')))

# Navigate to the About section
driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
about_section = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.XPATH, '//a[@data-control-name="contact_see_more"]')))
about_section.click()

# Wait for login fields to properly load
email_elem = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.NAME, 'session_key')))
password_elem = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.NAME, 'session_password')))
sign_in_button = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.XPATH, '//*[contains(@class, "sign-in-form__submit-button")]')))

# Login
email_elem.send_keys(username)
password_elem.send_keys(password)
sign_in_button.click()

# Wait for the job search bar to load
job_search_bar = WebDriverWait(driver, 20).until(
    EC.presence_of_element_located((By.XPATH, '//input[@aria-label="Search by title, skill, or company"]')))
job_search_bar.send_keys('software engineer')
job_search_bar.send_keys(Keys.RETURN)

# Retrieve the network log
performance_log = driver.get_log('performance')
for log in performance_log:
    message = json.loads(log['message'])['message']
    if message['method'] == 'Network.responseReceived':
        print(message['params']['response']['url'])

# Create a new document
document = Document()

# Retrieve Full Name, Headline, and Location
soup = BeautifulSoup(driver.page_source, 'lxml')
full_name = soup.find(
    'li', {'class': 'inline t-24 t-black t-normal break-words'}).get_text().strip()
headline = soup.find(
    'h2', {'class': 'mt1 t-18 t-black t-normal break-words'}).get_text().strip()
location = soup.find(
    'li', {'class': 't-16 t-black t-normal inline-block'}).get_text().strip()

# Add title and name
document.add_heading('Curriculum Vitae', 0)
document.add_heading(full_name, level=1)

# Add Full Name, Headline, and Location
document.add_heading(full_name, level=1)
document.add_paragraph(headline)
document.add_paragraph(location)

# Retrieve Work Experience
work_experience = []
for experience in soup.find_all('section', {'class': 'experience-section'}):
    title = experience.find(
        'h3', {'class': 't-16 t-black t-bold'}).get_text().strip()
    company = experience.find(
        'p', {'class': 'pv-entity__secondary-title t-14 t-black t-normal'}).get_text().strip()
    dates = experience.find('h4', {
                            'class': 'pv-entity__date-range t-14 t-black--light t-normal'}).get_text().strip()
    description = experience.find(
        'div', {'class': 'pv-entity__extra-details t-14 t-black t-normal'}).get_text().strip()
    work_experience.append(
        {'title': title, 'company': company, 'dates': dates, 'description': description})

# Add Work Experience
document.add_heading('Work Experience', level=1)
for item in work_experience:
    document.add_paragraph(
        f"{item['title']} at {item['company']}, {item['dates']}", style='List Bullet')
    document.add_paragraph(item['description'])
    document.add_paragraph('\n')

# Retrieve Education
education = []
for school in soup.find_all('section', {'class': 'education-section'}):
    degree = school.find(
        'h3', {'class': 't-16 t-black t-bold'}).get_text().strip()
    institution = school.find(
        'p', {'class': 'pv-entity__secondary-title t-14 t-black t-normal'}).get_text().strip()
    dates = school.find(
        'h4', {'class': 'pv-entity__dates t-14 t-black--light t-normal'}).get_text().strip()
    description = school.find('div', {
                              'class': 'pv-entity__extra-details t-14 t-black t-normal'}).get_text().strip()
    education.append({'degree': degree, 'institution': institution,
                     'dates': dates, 'description': description})

# Add Education
document.add_heading('Education', level=1)
for item in education:
    document.add_paragraph(
        f"{item['degree']} at {item['institution']}, {item['dates']}", style='List Bullet')
    document.add_paragraph(item['description'])
    document.add_paragraph('\n')

# Retrieve Skills
skills = []
for skill in soup.find_all('span', {'class': re.compile('^pv-skill-category-entity__name-text')}):
    skills.append(skill.get_text().strip())

# Add Skills
document.add_heading('Skills', level=1)
document.add_paragraph(', '.join(skills))

# Save the document
document.save('curriculum_vitae.docx')

# Quit the browser
driver.quit()
